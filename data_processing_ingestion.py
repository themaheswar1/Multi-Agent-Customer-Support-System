import os
import csv
import json
import hashlib
import pickle
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Any

import pdfplumber
import docx
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

load_dotenv()


# ----- Configuring the details -----

KNOWLEDGE_BASE_DIR = "knowledge_base"
FAISS_INDEX_PATH = "vectorstore/index.faiss"
METADATA_PATH = "vectorstore/metadata.pkl"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
CHUNK_SIZE = 300
CHUNK_OVERLAP = 50

@dataclass
class Chunk:
    chunk_id : str
    doc_id : str
    source : str
    file_type : str
    page : int
    chunk_index : int
    text : str
    metadata : Dict[str, Any] = field(default_factory=dict)

    def to_metadata(self) -> Dict[str, Any]:
        return {
            "chunk_id" : self.chunk_id,
            "doc_id" : self.doc_id,
            "source" : self.source,
            "file_type" : self.file_type,
            "page" : self.page,
            "chunk_index" : self.chunk_index,
            **self.metadata
        }
    
# --- Helpers --- 
def make_id(text: str) -> str:
    return hashlib.md5(text.encode()).hexdigest()[:12]

def chunk_text(text: str, source: str, file_type: str, page: int=0, extra_meta: Dict = {}) -> List[Chunk]:
    words = text.split()
    chunks = []
    i = 0
    idx = 0
    doc_id = make_id(source)   

    while i < len(words):
        chunk_words = words[i: i+CHUNK_SIZE] 
        chunk_text = " ".join(chunk_words).strip()
        if len(chunk_text) < 30:
            i += CHUNK_SIZE - CHUNK_OVERLAP
            continue

        chunk_id = make_id(f"{source}-{page}-{idx}")
        chunks.append(Chunk(
            chunk_id= chunk_id,
            doc_id=doc_id,
            source=source,
            file_type=file_type,
            page=page,
            chunk_index=idx,
            text=chunk_text,
            metadata=extra_meta.copy()
        ))
        idx += 1
        i += CHUNK_SIZE - CHUNK_OVERLAP 

    return chunks

# ---- doc loaders ----
def load_pdf(filepath: Path) -> List[Chunk]:
    chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if not text:
                continue
            text = " ".join(text.split())
            chunks.extend(chunk_text(
                text, filepath.name, "pdf", page=page_num
            ))
    print(f" [PDF]  {filepath.name} : {len(chunks)} chunks")
    return chunks   

def load_docx(filepath: Path) -> List[Chunk]:
    doc = docx.Document(str(filepath))
    chunks = []
    buffer = []
    page = 1

    for para in doc.paragraphs:
        try:
            text = para.text.strip() if para.text else ""
        except Exception:
            continue
        if not text:
            continue
        buffer.append(text)
        if len(buffer) >= 50:
            full_text = " ".join(buffer)
            chunks.extend(chunk_text(
                full_text, filepath.name, "docx",page=page
            ))
            buffer=[]
            page += 1
    if buffer:
        full_text = " ".join(buffer)
        chunks.extend(chunk_text(
            full_text, filepath.name, "docx", page=page
        ))

    for table in doc.tables:
        for row in table.rows:
            try:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text and cell.text.strip()
            )
                if row_text:
                    chunks.extend(chunk_text(
                        row_text, filepath.name, "docx", page=0,
                        extra_meta={"content_type": "table_row"}
                    ))
            except Exception as row_error:
                print(f"    Row error: {row_error}")
                continue  
    print(f"  [DOCX] {filepath.name}: {len(chunks)} chunks")
    return chunks              


def load_csv(filepath: Path) -> List[Chunk]:
    chunks = []
    with open(filepath, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        headers = reader.fieldnames or []

        for row_num, row in enumerate(reader, start=1):
            row_text= "  |  ".join(
                f"{k}:{v}" for k,v in row.items() if v and v.strip()
            )  
            if not row_text:
                continue
            chunk_id = make_id(f"{filepath.name}-{row_num}")
            doc_id = make_id(filepath.name)
            chunks.append(Chunk(
                chunk_id    = chunk_id,
                doc_id      = doc_id,
                source      = filepath.name,
                file_type   = "csv",
                page        = row_num,
                chunk_index = row_num,
                text        = row_text,
                metadata    = {"row_number": row_num, "columns": str(headers)}
            ))
    print(f"  [CSV]  {filepath.name}: {len(chunks)} chunks")
    return chunks

def load_txt(filepath : Path) -> List[Chunk]:
    text   = filepath.read_text(encoding='utf-8')
    text   = " ".join(text.split())
    chunks = chunk_text(text, filepath.name, "txt", page=1)
    print(f"  [TXT]  {filepath.name}: {len(chunks)} chunks")
    return chunks


# --- Main PipeLine ---
def load_all_documents() -> List[Chunk]:
    kb_path = Path(KNOWLEDGE_BASE_DIR)
    all_chunks : List[Chunk] = []

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".csv": load_csv,
        ".txt": load_txt,
    }

    files = sorted(kb_path.iterdir())
    print(f"\nFound {len(files)} files in {KNOWLEDGE_BASE_DIR}/\n")

    for filepath in files:
        ext = filepath.suffix.lower() #ext: extension
        if ext in loaders:
            try:
                chunks = loaders[ext](filepath)
                all_chunks.extend(chunks)
            except Exception as e:
                print(f" ---Error loading {filepath.name} : {e}")
        else:
            print(f" [SKIP] {filepath.name} (unsupported format)") 
    return all_chunks

def build_faiss_index(chunks: List[Chunk]):
    print(f"\n Embedding {len(chunks)} chunks using {EMBEDDING_MODEL}...")  
    model = SentenceTransformer(EMBEDDING_MODEL)
    texts = [c.text for c in chunks]
    embeddings = model.encode(texts, batch_size = 64, show_progress_bar=True, convert_to_numpy=True)

    #L2-normalize for cosine similarity
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
    embeddings = embeddings/np.maximum(norms, 1e-10)

    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings.astype(np.float32))

    faiss.write_index(index, FAISS_INDEX_PATH)

    metadata = [c.to_metadata() | {"text": c.text} for c in chunks]
    with open(METADATA_PATH,"wb") as f:
        pickle.dump(metadata,f)

    print(f"\n✅ FAISS index saved  → {FAISS_INDEX_PATH}")
    print(f"✅ Metadata saved     → {METADATA_PATH}")
    print(f"📊 Total vectors      : {index.ntotal}")
    print(f"📐 Embedding dim      : {dim}")


if __name__ == "__main__":
    print("=" * 55)
    print("  ShopSmart Knowledge Base — Ingestion Pipeline")
    print("=" * 55)

    chunks = load_all_documents()
    print(f"\nTotal chunks across all documents: {len(chunks)}")

    build_faiss_index(chunks)
    print("\nIngestion completes")


    


        


